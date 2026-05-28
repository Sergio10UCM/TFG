# Instalación de dependencias necesarias en el entorno virtual (ej. Google Colab)
# 1. pygam (versión de desarrollo compatible con Python 3.12)
# 2. shap y xgboost para explicabilidad y modelado no lineal
# 3. python-docx para la exportación de tablas en formato APA
!pip install git+https://github.com/dswah/pygam.git shap xgboost python-docx

# =============================================================================
# TFG — Capítulo 4: Análisis Inferencial (Parte Python)
# Cobertura: H4 (nacional vs internacional), H6 (evento × festivo + RF + GAM),
#            H7 (clustering + validación externa)
#
# PREREQUISITO: tfg_cap4_R.R ejecutado (genera los CSVs en cap4_outputs/)
#
# AUTOR: Sergio Díez Cardo — TFG Ciencia de Datos Aplicada
# =============================================================================

# Importación de librerías base y de tratamiento de datos
import os
import warnings
import numpy as np
import pandas as pd

# Modelado estadístico tradicional
import statsmodels.api as sm
import statsmodels.formula.api as smf

# Machine Learning: Clustering y preprocesamiento
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import silhouette_score

# Machine Learning: Modelos basados en árboles y métricas de importancia
from sklearn.ensemble import RandomForestRegressor
from sklearn.inspection import permutation_importance
import xgboost as xgb
import shap

# Modelos Aditivos Generalizados (GAM)
from pygam import LinearGAM, s, f

# Visualización
import matplotlib.pyplot as plt

# Silenciamiento de advertencias de versiones futuras para limpieza del output
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=UserWarning)

# Configuración del directorio de salida (se crea si no existe)
OUT_DIR = "cap4_outputs"
os.makedirs(OUT_DIR, exist_ok=True)

print("Paquetes cargados.\n")


# =============================================================================
# H4 — DIFERENCIACIÓN NACIONAL vs INTERNACIONAL
# =============================================================================
# FORMULACIÓN FORMAL:
#   H_0: β_internacional = β_nacional  (efecto idéntico sobre ambos)
#   H_1: β_internacional > β_nacional  (efecto mayor sobre internacionales)
#
# VARIABLES DEPENDIENTES: log(viajeros_nacionales), log(viajeros_internacionales)
# REGRESOR PRINCIPAL: log(n_conciertos_mes) ponderado por popularidad
#
# CONTRASTE: bootstrap no paramétrico de la diferencia de coeficientes
# =============================================================================

print("=" * 75)
print("H4 — TURISMO NACIONAL vs INTERNACIONAL")
print("=" * 75 + "\n")

# Carga y limpieza de datos mensuales exportados desde R
df_h4 = pd.read_csv(f"{OUT_DIR}/datos_h4.csv", parse_dates=["fecha_mes"])
df_h4 = df_h4.dropna(subset=["viajeros_espana", "viajeros_extranjero",
                              "n_conciertos_mes"]).copy()

# Transformaciones logarítmicas (elasticidades) para estabilizar varianza
df_h4["log_nac"]  = np.log(df_h4["viajeros_espana"])
df_h4["log_int"]  = np.log(df_h4["viajeros_extranjero"])
df_h4["log_conc"] = np.log(df_h4["n_conciertos_mes"] + 1)
df_h4["log_pop"]  = np.log(df_h4["max_listeners_mes"].fillna(1) + 1)

# Extracción de estacionalidad temporal
df_h4["mes"]      = df_h4["fecha_mes"].dt.month
df_h4["anio"]     = df_h4["fecha_mes"].dt.year

print(f"[H4.1] Observaciones mensuales (2022-2026): {len(df_h4)}\n")

# -----------------------------------------------------------------------------
# H4.2 — OLS paralelos
# -----------------------------------------------------------------------------

# Especificación base controlando por popularidad, mes (como categórica) y año
formula_base = "log_conc + log_pop + C(mes) + anio"

# Ajuste de Mínimos Cuadrados Ordinarios (OLS) para ambos subconjuntos
mod_nac = smf.ols(f"log_nac ~ {formula_base}", data=df_h4).fit()
mod_int = smf.ols(f"log_int ~ {formula_base}", data=df_h4).fit()

print("--- MODELO NACIONALES ---")
print(mod_nac.summary().tables[1])
print(f"R²: {mod_nac.rsquared:.4f}  |  R² ajustado: {mod_nac.rsquared_adj:.4f}\n")

print("--- MODELO INTERNACIONALES ---")
print(mod_int.summary().tables[1])
print(f"R²: {mod_int.rsquared:.4f}  |  R² ajustado: {mod_int.rsquared_adj:.4f}\n")

# Extracción de coeficientes y p-valores de la variable de interés
beta_nac = mod_nac.params["log_conc"]
beta_int = mod_int.params["log_conc"]
p_nac    = mod_nac.pvalues["log_conc"]
p_int    = mod_int.pvalues["log_conc"]

print(f"β(log_conc → log_nac): {beta_nac:+.4f}   p = {p_nac:.4f}")
print(f"β(log_conc → log_int): {beta_int:+.4f}   p = {p_int:.4f}")
print(f"Diferencia bruta: Δβ = {beta_int - beta_nac:+.4f}\n")

# -----------------------------------------------------------------------------
# H4.2bis — Test ANOVA F entre modelos anidados (lo pide el tutor)
# -----------------------------------------------------------------------------

print("[H4.2bis] Test ANOVA F: ¿aporta log(conc) información significativa?\n")

# Modelos restringidos (sin la variable de conciertos) para contrastar
mod_nac_restringido = smf.ols("log_nac ~ log_pop + C(mes) + anio", data=df_h4).fit()
mod_int_restringido = smf.ols("log_int ~ log_pop + C(mes) + anio", data=df_h4).fit()

print("--- NACIONALES: con vs sin log(conc) ---")
print(sm.stats.anova_lm(mod_nac_restringido, mod_nac))

print("\n--- INTERNACIONALES: con vs sin log(conc) ---")
print(sm.stats.anova_lm(mod_int_restringido, mod_int))

print("\n  Si F y p son significativos en ambos, la variable log(conc) aporta")
print("  información estadísticamente relevante en los dos modelos.\n")

# -----------------------------------------------------------------------------
# H4.3 — Bootstrap no paramétrico de la diferencia de coeficientes
# -----------------------------------------------------------------------------
# Necesitamos un IC para la diferencia. Como cada coef. procede de una OLS
# distinta sobre los mismos datos (correlación entre residuos), el bootstrap
# remuestreo no paramétrico es la solución más limpia.
# -----------------------------------------------------------------------------

print("[H4.3] Bootstrap no paramétrico (5000 réplicas)...\n")

rng = np.random.default_rng(42) # Semilla de reproducibilidad
B = 5000                        # Número de iteraciones
diff_boot = np.empty(B)
b_nac_boot = np.empty(B)
b_int_boot = np.empty(B)

n = len(df_h4)
# Bucle de remuestreo con reemplazo
for b in range(B):
    idx = rng.integers(0, n, n)
    sample = df_h4.iloc[idx]
    try:
        # Re-estimación de modelos en la muestra bootstrap
        mn = smf.ols(f"log_nac ~ {formula_base}", data=sample).fit()
        mi = smf.ols(f"log_int ~ {formula_base}", data=sample).fit()
        b_nac_boot[b] = mn.params["log_conc"]
        b_int_boot[b] = mi.params["log_conc"]
        diff_boot[b]  = b_int_boot[b] - b_nac_boot[b]
    except Exception:
        diff_boot[b] = np.nan

# Limpieza de nulos e inferencia a partir de la distribución empírica
diff_boot  = diff_boot[~np.isnan(diff_boot)]
ic_low     = np.percentile(diff_boot, 2.5)  # Límite inferior IC 95%
ic_high    = np.percentile(diff_boot, 97.5) # Límite superior IC 95%
# Cálculo del p-valor empírico bilateral
p_boot     = 2 * min(np.mean(diff_boot <= 0), np.mean(diff_boot >= 0))
mean_diff  = np.mean(diff_boot)

print(f"Δβ (media bootstrap): {mean_diff:+.4f}")
print(f"IC 95%: [{ic_low:+.4f}, {ic_high:+.4f}]")
print(f"p-value bootstrap (Δβ ≠ 0): {p_boot:.4f}\n")

# Visualización: Histograma del bootstrap
fig, ax = plt.subplots(figsize=(8, 5))
ax.hist(diff_boot, bins=50, edgecolor="white", color="#457B9D")
ax.axvline(0, color="black", linestyle="--", label="Δβ = 0")
ax.axvline(mean_diff, color="#E63946", linewidth=2, label=f"Media = {mean_diff:.3f}")
ax.axvline(ic_low, color="grey", linestyle=":", label="IC 95%")
ax.axvline(ic_high, color="grey", linestyle=":")
ax.set_xlabel("Δβ = β_internacional − β_nacional")
ax.set_ylabel("Frecuencia (5000 réplicas)")
ax.set_title("H4 — Distribución bootstrap de la diferencia de coeficientes")
ax.legend()
plt.tight_layout()
plt.savefig(f"{OUT_DIR}/h4_bootstrap_diff.png", dpi=180)
plt.close()


# -----------------------------------------------------------------------------
# INTERPRETACIÓN AUTOMÁTICA H4
# -----------------------------------------------------------------------------

print("-" * 75)
print("INTERPRETACIÓN AUTOMÁTICA — H4")
print("-" * 75 + "\n")

print("→ LÓGICA DE INTERPRETACIÓN:\n")
print("  H_0: β_internacional = β_nacional")
print("  H_1: β_internacional > β_nacional\n")

if p_boot < 0.05:
    print(f"  • p-value bootstrap = {p_boot:.4f} < 0.05 → RECHAZAMOS H_0")
    if mean_diff > 0:
        print(f"  • Δβ = {mean_diff:+.4f} > 0 con IC 95% [{ic_low:+.4f}, {ic_high:+.4f}]")
        print("  • El efecto de los conciertos sobre los viajeros INTERNACIONALES")
        print("    es significativamente mayor que sobre los nacionales.")
        print("  • La evidencia APOYA H_1\n")
    else:
        print(f"  • Δβ = {mean_diff:+.4f} < 0 → resultado contrario al esperado\n")
else:
    print(f"  • p-value bootstrap = {p_boot:.4f} ≥ 0.05 → NO rechazamos H_0")
    print("  • No hay evidencia significativa de diferencia entre efectos\n")

print(f"→ R² del modelo nacional:      {mod_nac.rsquared:.4f}")
print(f"→ R² del modelo internacional: {mod_int.rsquared:.4f}")
print("\n")


# =============================================================================
# H6 — INTERACCIÓN EVENTO × FESTIVO + RANDOM FOREST + GAM
# =============================================================================
# FORMULACIÓN FORMAL:
#   H_0: β_interaccion = 0   (el efecto del evento no se amplifica en festivos)
#   H_1: β_interaccion > 0   (festivos amplifican el efecto del evento)
#
# VARIABLE DEPENDIENTE: tasa_ocupacion diaria de Airbnb
#
# TRES APROXIMACIONES COMPLEMENTARIAS:
#   A) OLS con interacción explícita evento × festivo
#   B) Random Forest + importancia de variables + SHAP (no lineal)
#   C) GAM con suavizadores spline para los regresores numéricos
# =============================================================================

print("=" * 75)
print("H6 — INTERACCIÓN EVENTO × FESTIVO")
print("=" * 75 + "\n")

# Carga de la tabla diaria maestra
df_h6 = pd.read_csv(f"{OUT_DIR}/datos_h6_diario.csv", parse_dates=["fecha"])
df_h6 = df_h6.dropna(subset=["tasa_ocupacion"]).copy()

# Generación de variables booleanas y logarítmicas para el modelo
df_h6["es_gran_ev"]    = (df_h6["n_gran_evento"] > 0).astype(int)
df_h6["es_festivo"]    = df_h6["es_festivo"].astype(int)
df_h6["es_puente"]     = df_h6["es_puente"].astype(int)
df_h6["es_finde"]      = df_h6["es_finde"].astype(int)
df_h6["log_listeners"] = np.log10(df_h6["max_listeners"].fillna(1) + 1)
df_h6["log_conc"]      = np.log(df_h6["n_conciertos"] + 1)

print(f"[H6.1] Observaciones diarias con ocupación Airbnb: {len(df_h6)}\n")


# -----------------------------------------------------------------------------
# H6.A — OLS con interacción
# -----------------------------------------------------------------------------

print("[H6.A] OLS con interacción evento × festivo...\n")

# Ajuste lineal con término multiplicativo explícito (es_gran_ev * es_festivo)
mod_h6_ols = smf.ols(
    "tasa_ocupacion ~ es_gran_ev * es_festivo + log_conc + log_listeners "
    "+ es_finde + es_puente + C(dia_semana) + C(mes) + C(anio)",
    data=df_h6
).fit()

print(mod_h6_ols.summary().tables[1])
print(f"\n  R²: {mod_h6_ols.rsquared:.4f}  |  R² ajustado: {mod_h6_ols.rsquared_adj:.4f}\n")

beta_inter_ols = mod_h6_ols.params.get("es_gran_ev:es_festivo", np.nan)
p_inter_ols    = mod_h6_ols.pvalues.get("es_gran_ev:es_festivo", np.nan)


# -----------------------------------------------------------------------------
# H6.B — Random Forest + importancia de variables
# -----------------------------------------------------------------------------

print("[H6.B] Random Forest e importancia de variables...\n")

# Variable de interacción explícita materializada como feature (lo pide el tutor)
df_h6["interaccion_ev_fest"] = df_h6["es_gran_ev"] * df_h6["es_festivo"]

feat_cols = ["es_gran_ev", "log_conc", "log_listeners",
             "es_festivo", "es_puente", "es_finde",
             "mes", "anio",
             "interaccion_ev_fest"]   # interacción explícita

# Codificamos día de la semana como ordinal (lunes=1 … domingo=7) para el árbol
dia_map = {"Lunes":1,"Martes":2,"Miércoles":3,"Jueves":4,
           "Viernes":5,"Sábado":6,"Domingo":7}
df_h6["dia_num"] = df_h6["dia_semana"].map(dia_map)
feat_cols.append("dia_num")

# Extracción de matrices numpy para scikit-learn
X = df_h6[feat_cols].fillna(0).values
y = df_h6["tasa_ocupacion"].values

# Entrenamiento de Random Forest Regressor
rf = RandomForestRegressor(n_estimators=500, max_depth=10,
                           min_samples_leaf=5, random_state=42, n_jobs=-1)
rf.fit(X, y)

# Importancia por permutación (más fiable que Gini al evitar sesgo en variables de alta cardinalidad)
perm = permutation_importance(rf, X, y, n_repeats=10, random_state=42, n_jobs=-1)
imp_df = (pd.DataFrame({"feature": feat_cols,
                         "importance": perm.importances_mean,
                         "std": perm.importances_std})
          .sort_values("importance", ascending=False))

print("Importancia por permutación (Random Forest):")
print(imp_df.to_string(index=False))

# Visualización de las importancias calculadas
fig, ax = plt.subplots(figsize=(8, 5))
ax.barh(imp_df["feature"][::-1], imp_df["importance"][::-1],
        xerr=imp_df["std"][::-1], color="#F4A261", edgecolor="black")
ax.set_xlabel("Importancia por permutación")
ax.set_title("H6 — Importancia de variables (Random Forest)")
plt.tight_layout()
plt.savefig(f"{OUT_DIR}/h6_rf_importancia.png", dpi=180)
plt.close()

print(f"\n  R² in-sample del Random Forest: {rf.score(X, y):.4f}\n")


# -----------------------------------------------------------------------------
# H6.B' — XGBoost con SHAP para inspeccionar la interacción
# -----------------------------------------------------------------------------

print("[H6.B'] XGBoost + SHAP para la interacción evento × festivo...\n")

# Entrenamiento de Extreme Gradient Boosting para captura de no-linealidades complejas
xgb_mod = xgb.XGBRegressor(
    n_estimators=400, max_depth=6, learning_rate=0.05,
    subsample=0.85, colsample_bytree=0.85, random_state=42,
    n_jobs=-1, verbosity=0
)
xgb_mod.fit(X, y)
print(f"  R² in-sample XGBoost: {xgb_mod.score(X, y):.4f}\n")

# Extracción de Valores SHAP (SHapley Additive exPlanations)
explainer = shap.TreeExplainer(xgb_mod)
shap_values = explainer.shap_values(X)

# Análisis descriptivo: Efecto medio condicionado
mask_ev_fest    = (df_h6["es_gran_ev"] == 1) & (df_h6["es_festivo"] == 1)
mask_ev_no_fest = (df_h6["es_gran_ev"] == 1) & (df_h6["es_festivo"] == 0)
mask_no_ev      = (df_h6["es_gran_ev"] == 0)

mean_pred_ef = df_h6.loc[mask_ev_fest,    "tasa_ocupacion"].mean()
mean_pred_e  = df_h6.loc[mask_ev_no_fest, "tasa_ocupacion"].mean()
mean_pred_b  = df_h6.loc[mask_no_ev,      "tasa_ocupacion"].mean()

diff_ef_e = mean_pred_ef - mean_pred_e

print(f"  Ocupación media día evento + festivo: {mean_pred_ef:.4f}")
print(f"  Ocupación media día evento sin festivo: {mean_pred_e:.4f}")
print(f"  Ocupación media día sin evento:       {mean_pred_b:.4f}")
print(f"  Diferencia evento+festivo vs evento solo: {diff_ef_e:+.4f}\n")

# Generación del plot resumen de SHAP
fig = plt.figure(figsize=(8, 5))
shap.summary_plot(shap_values, X, feature_names=feat_cols, show=False)
plt.tight_layout()
plt.savefig(f"{OUT_DIR}/h6_shap_summary.png", dpi=180)
plt.close()


# -----------------------------------------------------------------------------
# H6.C — GAM con suavizadores (lo que el tutor llamó "GAN")
# -----------------------------------------------------------------------------
# Generalized Additive Models:
#   y = α + s(log_conc) + s(log_listeners) + f(dia_semana) + f(festivo) +
#       f(es_gran_ev) + interacción_suavizada(es_gran_ev × es_festivo)
#
# s() = suavizador spline; f() = factor (categórico)
# -----------------------------------------------------------------------------

print("[H6.C] GAM con suavizadores spline (incluye interacción explícita)...\n")

X_gam = df_h6[["es_gran_ev", "log_conc", "log_listeners",
               "es_festivo", "es_finde", "dia_num", "mes",
               "interaccion_ev_fest"]].fillna(0).values
y_gam = df_h6["tasa_ocupacion"].values

# Definición funcional del Modelo Aditivo Generalizado
gam = LinearGAM(
    f(0) +              # es_gran_ev (categórica)
    s(1) +              # log_conc (spline continuo)
    s(2) +              # log_listeners (spline continuo)
    f(3) +              # es_festivo (categórica)
    f(4) +              # es_finde (categórica)
    s(5, n_splines=7) + # dia_num (spline cíclico/ordinal)
    s(6, n_splines=10)+ # mes (spline estacional)
    f(7)                # interaccion_ev_fest (categórica, NUEVA)
).fit(X_gam, y_gam)

print(gam.summary())
print(f"\n  R² in-sample GAM: {gam.statistics_['pseudo_r2']['explained_deviance']:.4f}")
print(f"  AIC GAM: {gam.statistics_['AIC']:.2f}\n")

# Mirar el p-value del término 7 (interacción) en el summary:
# si p > 0,05 confirma el veredicto nulo de H6 también desde el GAM con interacción explícita.

# Gráfico de dependencia parcial: Efecto aislado de variables continuas tras suavizar
fig, axes = plt.subplots(1, 2, figsize=(12, 4))
for i, feat_idx in enumerate([1, 2]):
    XX = gam.generate_X_grid(term=feat_idx)
    pdep, confi = gam.partial_dependence(term=feat_idx, X=XX, width=0.95)
    axes[i].plot(XX[:, feat_idx], pdep, color="#457B9D", linewidth=2)
    axes[i].fill_between(XX[:, feat_idx], confi[:, 0], confi[:, 1],
                          color="#457B9D", alpha=0.25)
    axes[i].set_title(f"GAM: efecto suavizado de {['log_conc','log_listeners'][i]}")
    axes[i].set_xlabel(["log(n_conciertos)", "log10(max_listeners)"][i])
    axes[i].set_ylabel("Efecto parcial sobre tasa_ocupacion")
plt.tight_layout()
plt.savefig(f"{OUT_DIR}/h6_gam_efectos.png", dpi=180)
plt.close()


# -----------------------------------------------------------------------------
# INTERPRETACIÓN AUTOMÁTICA H6
# -----------------------------------------------------------------------------

print("-" * 75)
print("INTERPRETACIÓN AUTOMÁTICA — H6")
print("-" * 75 + "\n")

print("→ EVIDENCIA EN TRES NIVELES:\n")

print("NIVEL 1 — OLS con interacción explícita:")
if not np.isnan(p_inter_ols):
    print(f"  β(evento × festivo) = {beta_inter_ols:+.5f}, p = {p_inter_ols:.4f}")
    if p_inter_ols < 0.05 and beta_inter_ols > 0:
        print("  ✓ La interacción es positiva y significativa → APOYA H_1\n")
    elif p_inter_ols < 0.05:
        print("  ⚠ Significativa pero con signo no esperado\n")
    else:
        print("  ✗ No significativa — el OLS no encuentra evidencia de amplificación\n")

print("NIVEL 2 — Random Forest (importancia):")
print(f"  Variables más relevantes según permutación:")
top3 = imp_df.head(3)["feature"].tolist()
print(f"  {', '.join(top3)}\n")

print("NIVEL 3 — Diferencia descriptiva en SHAP / medias:")
print(f"  Ocupación promedio:")
print(f"    Día evento + festivo: {mean_pred_ef:.4f}")
print(f"    Día evento sin fest.: {mean_pred_e:.4f}")
print(f"    Diferencia bruta: {diff_ef_e:+.4f}")
if diff_ef_e > 0:
    print("  ✓ Coincidir con festivo → mayor ocupación → DIRECCIÓN consistente con H_1\n")
else:
    print("  ✗ Festivo no amplifica el efecto descriptivamente\n")

print("→ CONCLUSIÓN H6: la magnitud del efecto es pequeña incluso si es positiva.")
print("  Considerar que H6 puede ser estadísticamente débil aunque la dirección")
print("  apunte al efecto esperado.\n")


# =============================================================================
# H7 — TIPOLOGÍA LATENTE DE LA ACTIVIDAD CONCERTÍSTICA
# =============================================================================
# FORMULACIÓN FORMAL:
#   H_0: la actividad concertística es un único proceso homogéneo
#   H_1: existen al menos dos segmentos estructuralmente distintos
#        (escena local vs. gran formato global), y los meses dominados por
#        cada segmento difieren significativamente en sus indicadores
#        turísticos hoteleros
#
# DOS PASOS:
#   A) KMeans sobre features de los conciertos individuales
#      Selección de k por silhouette (probamos k=2,3,4)
#   B) Agregación a nivel mensual y validación externa contra INE
# =============================================================================

print("=" * 75)
print("H7 — TIPOLOGÍA LATENTE DE LA ACTIVIDAD CONCERTÍSTICA")
print("=" * 75 + "\n")

df_h7 = pd.read_csv(f"{OUT_DIR}/datos_h7.csv", parse_dates=["fecha"])
print(f"[H7.1] Conciertos cargados: {len(df_h7)}\n")

# -----------------------------------------------------------------------------
# H7.A — Features para clustering
# -----------------------------------------------------------------------------

# Generación de variables continuas y normalizadas para el algoritmo de distancias
df_h7["log_aforo"]     = np.log10(df_h7["aforo"].fillna(100) + 1)
df_h7["log_listeners"] = np.log10(df_h7["listeners"].fillna(1) + 1)
df_h7["rank_streams_inv"] = np.where(df_h7["ranking_streams"].notna(),
                                      1 / (df_h7["ranking_streams"] + 1),
                                      0)
df_h7["en_chart"]     = df_h7["en_chart_spain"].astype(int)
df_h7["gran_ev"]      = df_h7["es_gran_evento"].astype(int)

feat_h7 = ["log_aforo", "log_listeners", "rank_streams_inv", "en_chart", "gran_ev"]
X_h7 = df_h7[feat_h7].fillna(0).values

# Estandarización de medias a 0 y varianzas a 1 (obligatorio para K-Means)
scaler = StandardScaler()
X_h7_std = scaler.fit_transform(X_h7)

print(f"  Features: {feat_h7}")
print(f"  Observaciones para clustering: {len(X_h7_std)}\n")


# -----------------------------------------------------------------------------
# H7.B — Selección de k por silhouette
# -----------------------------------------------------------------------------

print("[H7.2] Selección de k por silhouette...\n")

# Búsqueda de hiperparámetros iterando posibles K clústeres
sil_scores = {}
for k in [2, 3, 4]:
    km = KMeans(n_clusters=k, n_init=20, random_state=42)
    labels = km.fit_predict(X_h7_std)
    # Cálculo del coeficiente de silueta (mide cohesión y separación)
    s_score = silhouette_score(X_h7_std, labels, sample_size=min(5000, len(X_h7_std)),
                                random_state=42)
    sil_scores[k] = s_score
    print(f"  k = {k}: silhouette = {s_score:.4f}")

k_opt = max(sil_scores, key=sil_scores.get)
print(f"\n  k óptimo: {k_opt} (silhouette = {sil_scores[k_opt]:.4f})\n")


# -----------------------------------------------------------------------------
# H7.C — KMeans con k óptimo + perfil de clusters
# -----------------------------------------------------------------------------

# Entrenamiento del modelo definitivo
km_final = KMeans(n_clusters=k_opt, n_init=50, random_state=42)
df_h7["cluster"] = km_final.fit_predict(X_h7_std)

print("[H7.3] Perfil de los clusters:")
# Resumen de las métricas reales para interpretar cada clúster conceptualmente
perfil = (df_h7.groupby("cluster")
          .agg(n_obs        = ("cluster", "size"),
               aforo_med    = ("aforo", "median"),
               listeners_med= ("listeners", "median"),
               pct_gran_ev  = ("gran_ev", "mean"),
               pct_chart    = ("en_chart", "mean")))
print(perfil)
print()


# -----------------------------------------------------------------------------
# H7.D — Agregación mensual y validación externa contra INE
# -----------------------------------------------------------------------------

print("[H7.4] Agregación mensual y validación externa contra INE...\n")

# Fijación de las fechas al primer día del mes para el cruce
df_h7["fecha_mes"] = df_h7["fecha"].values.astype("datetime64[M]")

# Proporción de conciertos del cluster "global" por mes
# El cluster "global" se infiere algorítmicamente buscando el de mayor aforo mediano
cluster_global = perfil["aforo_med"].idxmax()
print(f"  Cluster identificado como 'global': {cluster_global}")
print(f"  Aforo mediano del cluster global: {perfil.loc[cluster_global,'aforo_med']:.0f}\n")

mensual_h7 = (df_h7.assign(es_global = (df_h7["cluster"] == cluster_global).astype(int))
              .groupby("fecha_mes")
              .agg(n_conciertos = ("fecha", "count"),
                   prop_global  = ("es_global", "mean"))
              .reset_index())

# Cruzamos con datos hoteleros del INE (vienen en datos_h4.csv)
df_h4_merge = pd.read_csv(f"{OUT_DIR}/datos_h4.csv", parse_dates=["fecha_mes"])

mensual_merge = pd.merge(mensual_h7, df_h4_merge, on="fecha_mes", how="inner")
print(f"  Meses con datos INE + clustering: {len(mensual_merge)}\n")

# Test de correlación: prop_global vs ocupación / viajeros / pernoctaciones
from scipy.stats import spearmanr

vars_externas = ["ocupacion_plazas", "viajeros_total",
                  "viajeros_extranjero", "pernoctaciones_total"]

print("Correlaciones de la proporción del cluster global con indicadores INE:\n")
for v in vars_externas:
    rho, pv = spearmanr(mensual_merge["prop_global"], mensual_merge[v])
    print(f"  prop_global × {v:22s}  ρ = {rho:+.3f}  p = {pv:.4f}")

# Test de diferencia de medias: meses con prop_global alto vs bajo
mediana_prop = mensual_merge["prop_global"].median()
mes_alto = mensual_merge[mensual_merge["prop_global"] >  mediana_prop]
mes_bajo = mensual_merge[mensual_merge["prop_global"] <= mediana_prop]

# Contraste de hipótesis no paramétrico de Mann-Whitney
from scipy.stats import mannwhitneyu

print(f"\n  Mediana de prop_global: {mediana_prop:.3f}")
print(f"  Meses alta intensidad global: {len(mes_alto)} | baja: {len(mes_bajo)}\n")

print("Mann-Whitney U: meses alta vs baja intensidad global\n")
for v in vars_externas:
    u, p = mannwhitneyu(mes_alto[v].dropna(), mes_bajo[v].dropna(),
                         alternative="greater")
    media_alto = mes_alto[v].mean()
    media_bajo = mes_bajo[v].mean()
    diff_pct   = (media_alto - media_bajo) / media_bajo * 100
    print(f"  {v:22s}  U = {u:.0f}, p = {p:.4f}, "
          f"Δ = {diff_pct:+.1f}%")


# -----------------------------------------------------------------------------
# INTERPRETACIÓN AUTOMÁTICA H7
# -----------------------------------------------------------------------------

print("\n" + "-" * 75)
print("INTERPRETACIÓN AUTOMÁTICA — H7")
print("-" * 75 + "\n")

print("→ PASO 1 (estructura interna): KMeans encontró k =", k_opt,
      "clusters con silhouette =", f"{sil_scores[k_opt]:.4f}")
if sil_scores[k_opt] > 0.25:
    print("  ✓ La estructura de clusters es ESTADÍSTICAMENTE clara")
    print("    (silhouette > 0.25 indica separación razonable)")
else:
    print("  ⚠ Silhouette modesto: clusters poco separados")
print()

print("→ PASO 2 (validación externa): contraste con INE")
print(f"  El cluster 'global' tiene aforo mediano de "
      f"{perfil.loc[cluster_global,'aforo_med']:.0f},")
print(f"  oyentes mediano de {perfil.loc[cluster_global,'listeners_med']:.0f}")
print(f"  y un {perfil.loc[cluster_global,'pct_gran_ev']*100:.1f}% son grandes eventos.\n")

print("→ Si los meses con alta proporción del cluster global tienen")
print("  significativamente más viajeros internacionales y más pernoctaciones,")
print("  entonces los clusters tienen RELEVANCIA TURÍSTICA externa, no solo")
print("  estadística interna. La evidencia conjunta APOYA H_1.\n")


# Exportar tabla resumen de clusters a CSV
perfil.to_csv(f"{OUT_DIR}/h7_perfil_clusters.csv")
mensual_merge.to_csv(f"{OUT_DIR}/h7_mensual_validacion.csv", index=False)


# =============================================================================
# RESUMEN FINAL
# =============================================================================

print("\n" + "=" * 75)
print("SCRIPT PYTHON COMPLETADO.")
print("=" * 75)
print("\nArchivos generados en cap4_outputs/:")
for f in sorted(os.listdir(OUT_DIR)):
    print(f"  {f}")
print()

# Enlazar con Google Drive si se ejecuta dentro de Colab
from google.colab import drive
drive.mount('/content/drive')

# Verificación robusta mediante Validación Cruzada K-Fold para XGBoost
from sklearn.model_selection import KFold, cross_val_score

cv = KFold(n_splits=5, shuffle=True, random_state=42)
cv_scores = cross_val_score(xgb_mod, X, y, cv=cv, scoring="r2", n_jobs=-1)
print(f"  R² 5-fold CV XGBoost: {cv_scores.mean():.4f} ± {cv_scores.std():.4f}")
print(f"  R² in-sample (sobreajustado): {xgb_mod.score(X, y):.4f}")

########

# Corrección analítica: Alineación de los tipos temporales para la unión correcta
df_h7["fecha_mes"] = pd.to_datetime(df_h7["fecha"]).dt.to_period("M").dt.to_timestamp()
df_h4_merge["fecha_mes"] = pd.to_datetime(df_h4_merge["fecha_mes"]).dt.to_period("M").dt.to_timestamp()

# Reforzamiento manual con k=2 (Re-clusterización explícita garantizada)
km_2 = KMeans(n_clusters=2, n_init=50, random_state=42)
df_h7["cluster_2"] = km_2.fit_predict(X_h7_std)

# Re-identificación algorítmica del cluster global tras forzar k=2
perfil_2 = df_h7.groupby("cluster_2").agg(
    aforo_med=("aforo", "median"),
    pct_gran=("gran_ev", "mean"))
cluster_global_2 = perfil_2["aforo_med"].idxmax()

mensual_h7 = (df_h7.assign(es_global=(df_h7["cluster_2"] == cluster_global_2).astype(int))
              .groupby("fecha_mes")
              .agg(n_conciertos=("fecha", "count"),
                   prop_global=("es_global", "mean"))
              .reset_index())

mensual_merge = pd.merge(mensual_h7, df_h4_merge, on="fecha_mes", how="inner")

# Comprobación de correlación final evadiendo problemas de valores NaN con máscaras
for v in vars_externas:
    mask = mensual_merge[["prop_global", v]].notna().all(axis=1)
    if mask.sum() >= 10:
        rho, pv = spearmanr(mensual_merge.loc[mask, "prop_global"],
                            mensual_merge.loc[mask, v])
        print(f"  prop_global × {v:22s}  ρ = {rho:+.3f}  p = {pv:.4f}  (n={mask.sum()})")

# =============================================================================
# CAPÍTULO 4 — TABLAS APA Y FIGURAS DESDE PYTHON
# Tablas: 4.6, 4.10, 4.11
# Figuras: 4.3
# =============================================================================

# Automatización de la generación de reporte de Word (DOCX)
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = "cap4_outputs/tablas_apa"
FIG_DIR = "cap4_outputs/figuras"
os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(FIG_DIR, exist_ok=True)

# -----------------------------------------------------------------------------
# Funciones auxiliares para formateo de salida
# -----------------------------------------------------------------------------

def fmt_es(x, dig=3):
    """Formato número estilo español: coma decimal y puntos de miles."""
    if pd.isna(x):
        return ""
    return f"{x:,.{dig}f}".replace(",", "X").replace(".", ",").replace("X", ".")

def fmt_p(p):
    """Formato de p-valor bajo convención estándar APA."""
    if pd.isna(p):
        return ""
    if p < 0.001:
        return "< 0,001"
    return fmt_es(p, 3)

def apa_save(df, nombre, numero, titulo, nota):
    """Genera un archivo .docx exportando el DataFrame en estricto formato APA 7."""
    doc = Document()

    # Estilo tipográfico base
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Inserción del Número de tabla (negrita)
    p_num = doc.add_paragraph()
    run = p_num.add_run(numero)
    run.bold = True
    run.font.size = Pt(11)

    # Inserción del Título (cursiva)
    p_tit = doc.add_paragraph()
    run = p_tit.add_run(titulo)
    run.italic = True
    run.font.size = Pt(11)

    # Creación de la estructura base de la Tabla
    tabla = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    tabla.style = "Table Grid"  # se ajustará después mediante código XML

    # Volcado de la Cabecera
    for j, col in enumerate(df.columns):
        cell = tabla.rows[0].cells[j]
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"

    # Volcado del cuerpo de los Datos
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = tabla.rows[i + 1].cells[j]
            cell.text = str(df.iat[i, j])
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"

    # Manipulación a bajo nivel (XML) para aplicar formato APA: solo líneas horizontales
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import OxmlElement, parse_xml

    def set_cell_border(cell, **kwargs):
        """Función auxiliar inyectada para remover bordes verticales"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            if border_name in kwargs:
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), str(kwargs[border_name]))
                border.set(qn("w:color"), "000000")
            else:
                border.set(qn("w:val"), "nil")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    n_rows = df.shape[0] + 1
    for j in range(df.shape[1]):
        set_cell_border(tabla.rows[0].cells[j], top=12, bottom=6)
        set_cell_border(tabla.rows[n_rows - 1].cells[j], bottom=12)
        for i in range(1, n_rows - 1):
            set_cell_border(tabla.rows[i].cells[j])

    # Volcado del pié de Nota
    doc.add_paragraph()
    p_nota = doc.add_paragraph()
    run_n = p_nota.add_run("Nota. ")
    run_n.italic = True
    run_n.font.size = Pt(10)
    run_t = p_nota.add_run(nota)
    run_t.font.size = Pt(10)

    # Persistencia del documento final
    ruta = os.path.join(OUT_DIR, f"{nombre}.docx")
    doc.save(ruta)
    print(f"✓ Guardada: {ruta}")


# =============================================================================
# TABLA 4.6 — Elasticidades H4 (nacional vs internacional)
# =============================================================================

df_4_6 = pd.DataFrame({
    "Modelo":          ["log(viajeros nacionales)", "log(viajeros internacionales)"],
    "β̂ log(conc)":    ["+" + fmt_es(0.0894, 4), "+" + fmt_es(0.2102, 4)],
    "SE":              [fmt_es(0.026, 3),         fmt_es(0.041, 3)],
    "p-valor":         [fmt_p(0.001),             fmt_p(0.0001)],
    "R²":              [fmt_es(0.799, 3),         fmt_es(0.930, 3)],
})
# y en la nota: Δβ̂ = +0,1295; IC 95% = [+0,029; +0,245]; p = 0,011 (5.000 réplicas).


apa_save(df_4_6, "tabla_4_6_elasticidades_h4",
    numero = "Tabla 4.6",
    titulo = "Elasticidades del turismo nacional e internacional respecto a la actividad concertística",
    nota   = ("Ambos modelos incluyen controles de mes, año y log(popularidad máxima). "
              "n = 51 meses (2022–2026). Bootstrap no paramétrico sobre Δβ = β_int − β_nac: "
              "Δβ̂ = +0,1713; IC 95% = [+0,053; +0,264]; p = 0,003 (5.000 réplicas).")
)


# =============================================================================
# TABLA 4.10 — Perfil de clusters K=2 (H7)
# =============================================================================

df_4_10 = pd.DataFrame({
    "Cluster":           ["C0 (escena local)", "C1 (gran formato global)"],
    "n":                 ["8.724",   "2.272"],
    "Aforo mediano":     ["640",     "16.000"],
    "Oyentes medianos":  ["n.d.",    "3.281.433"],
    "% gran evento":     ["0%",      "100%"],
    "% chart España":    [fmt_es(4.0, 1) + "%", fmt_es(41.1, 1) + "%"],
})

apa_save(df_4_10, "tabla_4_10_perfil_clusters_h7",
    numero = "Tabla 4.10",
    titulo = "Perfil de los clusters con K = 2",
    nota   = ("'n.d.' indica que la mayoría de artistas del cluster local no tienen datos "
              "de oyentes en el dataset de Kworb, coherente con su naturaleza de escena "
              "local de pequeño formato.")
)


# =============================================================================
# TABLA 4.11 — Correlaciones Spearman (H7)
# =============================================================================

df_4_11 = pd.DataFrame({
    "Indicador hotelero": ["Ocupación hotelera por plazas", "Viajeros totales",
                           "Viajeros internacionales", "Pernoctaciones totales"],
    "ρ":                  ["+" + fmt_es(0.235, 3), "+" + fmt_es(0.248, 3),
                           "+" + fmt_es(0.144, 3), "+" + fmt_es(0.221, 3)],
    "p-valor":            [fmt_p(0.097), fmt_p(0.079), fmt_p(0.312), fmt_p(0.118)],
})

apa_save(df_4_11, "tabla_4_11_spearman_h7",
    numero = "Tabla 4.11",
    titulo = "Correlaciones de Spearman entre la proporción mensual del cluster global y los indicadores hoteleros del INE",
    nota   = "n = 51 meses (2022–2026). Test bilateral."
)

# =============================================================================
# TABLA 5.6 — OLS con interacción evento × festivo (H6)
# =============================================================================

# Sacar coeficientes relevantes del mod_h6_ols ya estimado
coefs_relevantes = ["es_gran_ev", "es_festivo", "es_gran_ev:es_festivo",
                    "log_conc", "log_listeners", "es_finde", "es_puente"]

filas_h6 = []
for c in coefs_relevantes:
    if c in mod_h6_ols.params.index:
        b = mod_h6_ols.params[c]
        se = mod_h6_ols.bse[c]
        p  = mod_h6_ols.pvalues[c]
        # Cálculo de significancia en formato de asteriscos para APA
        estrellas = ("***" if p < 0.001 else
                     "**"  if p < 0.01  else
                     "*"   if p < 0.05  else "")
        etiqueta = {
            "es_gran_ev":            "Gran evento",
            "es_festivo":            "Festivo",
            "es_gran_ev:es_festivo": "Gran evento × Festivo",
            "log_conc":              "log(nº conciertos)",
            "log_listeners":         "log₁₀(listeners)",
            "es_finde":              "Fin de semana",
            "es_puente":             "Puente",
        }[c]
        filas_h6.append({
            "Variable": etiqueta,
            "β̂":       ("+" if b >= 0 else "") + fmt_es(b, 4) + estrellas,
            "SE":      fmt_es(se, 4),
            "p-valor": fmt_p(p),
        })

df_5_6 = pd.DataFrame(filas_h6)

apa_save(df_5_6, "tabla_5_6_ols_h6",
    numero = "Tabla 5.6",
    titulo = "OLS con interacción evento × festivo sobre la ocupación diaria agregada de Airbnb",
    nota   = (f"n = {int(mod_h6_ols.nobs)} días. Variable dependiente: tasa de ocupación diaria de Airbnb. "
              f"Controles adicionales no reportados: día de la semana, mes y año (variables categóricas). "
              f"R² = {mod_h6_ols.rsquared:.3f}; R² ajustado = {mod_h6_ols.rsquared_adj:.3f}. "
              f"* p < 0,05; ** p < 0,01; *** p < 0,001.")
)

# =============================================================================
# FIGURA 4.3 — Validación externa H7: dispersión mes a mes
# =============================================================================

mensual_h7 = pd.read_csv("cap4_outputs/h7_mensual_validacion.csv")

# Parametrización estricta de Matplotlib para adecuarse a estética APA (fuente serif)
plt.rcParams.update({
    "font.family": "serif",
    "font.serif":  ["Times New Roman"],
    "font.size":   11,
    "axes.titlesize":   13,
    "axes.titleweight": "bold",
    "axes.labelsize":   11,
    "axes.spines.top":   False,
    "axes.spines.right": False,
})

fig, ax = plt.subplots(figsize=(9, 5))

# Filtrar NaN
mask = mensual_h7[["prop_global", "ocupacion_plazas"]].notna().all(axis=1)
x = mensual_h7.loc[mask, "prop_global"].values
y = mensual_h7.loc[mask, "ocupacion_plazas"].values

# Scatter Plot de la dispersión
ax.scatter(x, y, color="#457B9D", s=50, alpha=0.75, edgecolor="white",
           linewidth=0.5)

# Línea de regresión principal calculada con SciPy
from scipy import stats as sps
slope, intercept, r_val, p_val, _ = sps.linregress(x, y)
x_line = np.linspace(x.min(), x.max(), 100)
y_line = intercept + slope * x_line
ax.plot(x_line, y_line, color="#E63946", linewidth=1.5)

# Cálculo manual de la Banda de Confianza al 95% para la regresión
n = len(x)
y_pred = intercept + slope * x
se = np.sqrt(np.sum((y - y_pred)**2) / (n - 2))
t_val = sps.t.ppf(0.975, n - 2)
ci = t_val * se * np.sqrt(1/n + (x_line - x.mean())**2 / np.sum((x - x.mean())**2))
ax.fill_between(x_line, y_line - ci, y_line + ci, color="#E63946", alpha=0.15)

ax.set_xlabel("Proporción mensual de conciertos del cluster global")
ax.set_ylabel("Tasa de ocupación hotelera (%)")
ax.text(0.97, 0.05, "ρ = +0,235   p = 0,097",
        transform=ax.transAxes, ha="right", va="bottom",
        fontsize=11, family="serif",
        bbox=dict(boxstyle="round,pad=0.4", facecolor="white",
                  edgecolor="grey", alpha=0.9))

# Formatear eje x como porcentaje
ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{int(x*100)}%"))

plt.suptitle("Figura 4.3", x=0.05, ha="left", fontsize=13, fontweight="bold")
plt.title("Validación externa de H7: proporción del cluster global frente a ocupación hotelera",
          loc="left", fontsize=11, color="grey", pad=20)
plt.figtext(0.05, 0.01, "n = 51 meses (2022–2026). Correlación de Spearman bilateral.",
            ha="left", fontsize=9, color="grey")

plt.tight_layout(rect=[0, 0.03, 1, 0.97])
plt.savefig(f"{FIG_DIR}/fig_4_3_validacion_h7.png",
            dpi=300, bbox_inches="tight", facecolor="white")
plt.close()
print(f"✓ Guardada: {FIG_DIR}/fig_4_3_validacion_h7.png")

# Puente para forzar la descarga de los ficheros al disco local desde Google Colab
from google.colab import files

# Ruta donde tu script guarda la tabla de la H6
ruta_h6 = "cap4_outputs/tablas_apa/tabla_5_6_ols_h6.docx"

if os.path.exists(ruta_h6):
    files.download(ruta_h6)
else:
    print("El archivo aún no se ha generado. Asegúrate de ejecutar la celda de la Tabla 5.6")